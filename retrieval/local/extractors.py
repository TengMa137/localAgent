from __future__ import annotations

import mimetypes
import re
from pathlib import Path
from typing import Any, Callable

from tools.filesystem.errors import ValidationError


Extractor = Callable[[Path], tuple[str, dict[str, Any]]]
_EXTRACTORS: dict[str, Extractor] = {}


def extractor(*suffixes: str) -> Callable[[Extractor], Extractor]:
    def _wrap(fn: Extractor) -> Extractor:
        for s in suffixes:
            _EXTRACTORS[s.lower()] = fn
        return fn
    return _wrap


def list_supported_suffixes() -> list[str]:
    return sorted(_EXTRACTORS.keys())


def guess_mime(path: Path) -> str:
    mime, _ = mimetypes.guess_type(path.name)
    return mime or "application/octet-stream"


def safe_decode_bytes(b: bytes) -> str:
    try:
        return b.decode("utf-8")
    except UnicodeDecodeError:
        return b.decode("latin-1", errors="replace")


def normalize_whitespace(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def extract_text(path: Path) -> tuple[str, dict[str, Any]]:
    """
    Extract text from a file based on suffix.
    Falls back to decoding bytes as text for unknown suffixes.
    """
    ext = path.suffix.lower()
    fn = _EXTRACTORS.get(ext)
    if fn is None:
        b = path.read_bytes()
        return safe_decode_bytes(b), {"note": f"no extractor for {ext}; decoded as text"}
    return fn(path)


# -------------------------
# Outline helpers
# -------------------------

def _outline_item(
    *,
    level: int,
    title: str,
    start: int | None = None,
    end: int | None = None,
    page: int | None = None,
) -> dict[str, Any]:
    t = (title or "").strip()
    if not t:
        return {}
    return {"level": int(level), "title": t, "start": start, "end": end, "page": page}


def _line_start_offsets(text: str) -> list[int]:
    # returns list of start indices for each line (0-based line number)
    starts = [0]
    for m in re.finditer(r"\n", text):
        starts.append(m.end())
    return starts


def _char_span_for_line(text: str, line_starts: list[int], line_no: int) -> tuple[int, int]:
    if line_no < 0:
        return (0, 0)
    if line_no >= len(line_starts):
        return (len(text), len(text))
    start = line_starts[line_no]
    end = line_starts[line_no + 1] - 1 if line_no + 1 < len(line_starts) else len(text)
    return (start, end)


@extractor(".txt", ".csv", ".log")
def _extract_text_like(path: Path) -> tuple[str, dict[str, Any]]:
    return safe_decode_bytes(path.read_bytes()), {}


@extractor(".md", ".markdown")
def _extract_markdown(path: Path) -> tuple[str, dict[str, Any]]:
    md = safe_decode_bytes(path.read_bytes())

    try:
        from markdown_it import MarkdownIt  # type: ignore
    except Exception as e:
        raise ValidationError(
            "Markdown outline extraction requires 'markdown-it-py'. Install it (pip install markdown-it-py)."
        ) from e

    # Parse tokens to find heading_open/inline/heading_close
    # Use token.map for line numbers when available.
    mi = MarkdownIt()
    tokens = mi.parse(md)

    line_starts = _line_start_offsets(md)
    outline: list[dict[str, Any]] = []

    i = 0
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == "heading_open":
            # heading tag like h1..h6
            try:
                level = int(tok.tag[1]) if tok.tag and tok.tag.startswith("h") else 1
            except Exception:
                level = 1

            # next token is usually inline with content
            title = ""
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                title = tokens[i + 1].content or ""

            # map gives [start_line, end_line]
            start = end = None
            if tok.map and len(tok.map) >= 1:
                start_line = tok.map[0]
                s, e = _char_span_for_line(md, line_starts, start_line)
                start, end = s, e

            item = _outline_item(level=level, title=title, start=start, end=end)
            if item:
                outline.append(item)

        i += 1

    meta: dict[str, Any] = {"format": "markdown"}
    if outline:
        meta["outline"] = outline
    return md, meta


@extractor(".json")
def _extract_json(path: Path) -> tuple[str, dict[str, Any]]:
    return safe_decode_bytes(path.read_bytes()), {"format": "json"}


@extractor(".yaml", ".yml")
def _extract_yaml(path: Path) -> tuple[str, dict[str, Any]]:
    return safe_decode_bytes(path.read_bytes()), {"format": "yaml"}


@extractor(".html", ".htm")
def _extract_html(path: Path) -> tuple[str, dict[str, Any]]:
    html = safe_decode_bytes(path.read_bytes())
    try:
        from bs4 import BeautifulSoup  # type: ignore

        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        title = (soup.title.string.strip() if soup.title and soup.title.string else "")

        # Build text while preserving heading markers to get reliable offsets.
        # We only include headings + paragraph-ish blocks; this keeps structure.
        blocks: list[tuple[str, int]] = []  # (block_text, heading_level(0 if none))
        outline: list[dict[str, Any]] = []

        # Simple traversal: take headings + paragraphs + list items + pre blocks in doc order
        for el in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre", "blockquote"]):
            txt = el.get_text(" ", strip=True)
            if not txt:
                continue
            if el.name and el.name.startswith("h"):
                lvl = int(el.name[1])
                blocks.append((txt, lvl))
            else:
                blocks.append((txt, 0))

        parts: list[str] = []
        pos = 0
        for txt, lvl in blocks:
            if lvl > 0:
                # heading line
                line = f"{'#' * lvl} {txt}\n"
                start = pos
                parts.append(line)
                pos += len(line)
                end = pos
                item = _outline_item(level=lvl, title=txt, start=start, end=end)
                if item:
                    outline.append(item)
                # blank line after heading to separate sections
                parts.append("\n")
                pos += 1
            else:
                line = txt + "\n"
                parts.append(line)
                pos += len(line)

        text = "".join(parts).strip()

        meta: dict[str, Any] = {}
        if title:
            meta["title"] = title
        if outline:
            meta["outline"] = outline
        meta["format"] = "html"
        return text, meta

    except Exception:
        # naive fallback (no outline)
        text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", html)
        text = re.sub(r"(?is)<.*?>", " ", text)
        text = re.sub(r"&nbsp;", " ", text)
        text = re.sub(r"&amp;", "&", text)
        text = re.sub(r"&lt;", "<", text)
        text = re.sub(r"&gt;", ">", text)
        return text, {"title": "", "format": "html"}


@extractor(".docx")
def _extract_docx(path: Path) -> tuple[str, dict[str, Any]]:
    from docx import Document  # type: ignore

    doc = Document(str(path))
    parts: list[str] = []
    outline: list[dict[str, Any]] = []

    pos = 0

    def add_line(s: str) -> tuple[int, int]:
        nonlocal pos
        line = s.rstrip() + "\n"
        start = pos
        parts.append(line)
        pos += len(line)
        end = pos
        return start, end

    # Paragraphs (use styles for headings)
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue

        style_name = ""
        try:
            style_name = (p.style.name or "").strip()
        except Exception:
            style_name = ""

        # Heading 1..9 commonly exist as styles "Heading 1", "Heading 2", etc.
        lvl = 0
        m = re.match(r"^Heading\s+(\d+)$", style_name, flags=re.IGNORECASE)
        if m:
            lvl = int(m.group(1))

        if lvl > 0:
            start, end = add_line(f"{'#' * min(6, lvl)} {txt}")
            item = _outline_item(level=min(6, lvl), title=txt, start=start, end=end)
            if item:
                outline.append(item)
            add_line("")  # blank line
        else:
            add_line(txt)

    # Tables: include as text (but not in outline)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                add_line("\t".join(cells))

    text = "".join(parts).strip()
    meta: dict[str, Any] = {"format": "docx"}
    if outline:
        meta["outline"] = outline
    return text, meta


@extractor(".pdf")
def _extract_pdf(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as e:
        print(
            "\n[ERROR] PDF extraction requires the 'pypdf' package.\n"
            "Install it with:\n"
            "  pip install pypdf\n",
        )
        raise ValidationError(
            "PDF extraction requires the 'pypdf' package. Install it with: pip install pypdf"
        ) from e

    
    reader = PdfReader(str(path))

    # ---- Extract outline/bookmarks best-effort ----
    outline_items: list[dict[str, Any]] = []
    try:
        raw_outline = None
        if hasattr(reader, "outlines"):
            raw_outline = reader.outlines
        elif hasattr(reader, "outline"):
            raw_outline = reader.outline  # type: ignore[attr-defined]

        def add_item(level: int, it: Any) -> None:
            title = getattr(it, "title", None) or getattr(it, "name", None) or ""
            page_idx = None
            try:
                if hasattr(reader, "get_destination_page_number"):
                    page_idx = reader.get_destination_page_number(it)  # type: ignore[arg-type]
            except Exception:
                page_idx = None

            item = _outline_item(
                level=level,
                title=str(title),
                page=(page_idx + 1) if isinstance(page_idx, int) else None,  # store 1-based
            )
            if item:
                outline_items.append(item)

        def walk_list(lst: list[Any], level: int) -> None:
            for it in lst:
                if isinstance(it, list):
                    walk_list(it, level + 1)
                else:
                    add_item(level, it)

        if isinstance(raw_outline, list) and raw_outline:
            walk_list(raw_outline, 1)

    except Exception:
        outline_items = []

    # ---- Extract page text + compute page_char_starts ----
    pages_text: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        pages_text.append(t.strip() or f"[Page {i+1}: no extractable text]")

    # Build joined text and page_char_starts precisely
    parts: list[str] = []
    page_char_starts: list[int] = []
    pos = 0
    for i, ptext in enumerate(pages_text):
        page_char_starts.append(pos)
        parts.append(ptext)
        pos += len(ptext)
        if i != len(pages_text) - 1:
            parts.append("\n\n")
            pos += 2

    text = "".join(parts)

    meta: dict[str, Any] = {
        "pages": len(reader.pages),
        "format": "pdf",
        "page_char_starts": page_char_starts,  # <-- new
    }
    if outline_items:
        meta["outline"] = outline_items

    return text, meta
